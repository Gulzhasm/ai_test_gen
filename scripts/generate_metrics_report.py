#!/usr/bin/env python3
"""Generate Phase 1 Empirical Metrics Report as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def create_report():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ===== TITLE PAGE =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phase 1 — Empirical Performance Metrics")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Driven Test Case Generation:\nA Hybrid Rule-Based + LLM Approach")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    case_study = doc.add_paragraph()
    case_study.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = case_study.add_run("Case Study: ENV QuickDraw Application")
    run.font.size = Pt(13)
    run.italic = True

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Data Source: Azure DevOps (cdpinc/Env)\n"
        f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n"
        "LLM Provider: Gemini 2.5 Flash\n"
        "Author: Gulzhas Mailybayeva"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS (manual) =====
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Research Questions",
        "3. Project Scale — Azure DevOps Data",
        "4. AI-Test-Gen Tool — Generation Metrics",
        "5. Complete User Story & Test Case Inventory",
        "6. Time & Productivity Analysis",
        "7. LLM Cost Analysis",
        "8. Quality Metrics",
        "9. Coverage Analysis",
        "10. Comparison with Related Work",
        "11. Key Performance Indicators (KPIs)",
        "12. Threats to Validity",
        "13. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents the empirical performance metrics collected during Phase 1 of the "
        "AI-Driven Test Case Generation project, a hybrid approach combining deterministic "
        "rule-based generation with Large Language Model (LLM) correction for producing "
        "manual functional test cases from user stories."
    )
    doc.add_paragraph(
        "The system was evaluated on the ENV QuickDraw desktop/tablet application, "
        "a CAD-style drawing tool under active development. All data was collected from "
        "the production Azure DevOps (ADO) instance and local generation output files."
    )

    doc.add_heading("Key Findings", level=2)
    key_findings = [
        ("55", "user stories on the ADO board across 3 development phases"),
        ("870", "test cases in Azure DevOps Test Plans (50 stories with test coverage)"),
        ("207", "test cases with full traceability data from 12 locally-generated stories"),
        ("1,563", "individual test steps generated across those 12 stories"),
        ("92%", "time reduction compared to fully manual test case writing"),
        ("$0.002", "LLM cost per test case (Gemini 2.5 Flash)"),
        ("72.9%", "first-pass quality rate (tests needing zero corrections)"),
        ("2.3x", "acceptance-criteria-to-test-case amplification ratio"),
    ]
    for val, desc in key_findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(val)
        run.bold = True
        p.add_run(f" — {desc}")

    doc.add_page_break()

    # ===== 2. RESEARCH QUESTIONS =====
    doc.add_heading("2. Research Questions", level=1)

    doc.add_paragraph(
        "This empirical evaluation is structured around four research questions that guide "
        "the assessment of the hybrid rule-based + LLM test generation approach."
    )

    rqs = [
        ("RQ1: ", "Does hybrid rule-based + LLM generation reduce manual test design time "
         "compared to fully manual methods?\n"
         "Metric: Time reduction (%), cost comparison at $40/hr conservative baseline [11]."),
        ("RQ2: ", "What is the first-pass structural quality rate of generated functional test cases?\n"
         "Metric: Percentage of test cases requiring zero human corrections across 6 quality "
         "dimensions (AC coverage, determinism, grounding, logic, consistency, uniqueness)."),
        ("RQ3: ", "What are the dominant failure modes in LLM-assisted test generation?\n"
         "Metric: Issue classification by type, severity, and frequency across manual review passes."),
        ("RQ4: ", "How does the hybrid approach compare to pure LLM approaches reported in "
         "prior literature?\n"
         "Metric: First-pass quality, AC coverage, cost per test case, and time reduction "
         "benchmarked against Li & Yuan (2024) [4], Masuda et al. (2025) [5], "
         "and Xue et al. (ISSTA 2024) [6]."),
    ]
    for prefix, desc in rqs:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()
    doc.add_heading("2.1 Research Question Mapping", level=2)
    add_styled_table(doc,
        ["Research Question", "Report Sections", "Phase 1 Finding"],
        [
            ["RQ1 — Time & cost reduction",
             "Sections 6, 7",
             "92% time reduction; $0.002/TC; 12.5x cheaper incl. review"],
            ["RQ2 — First-pass quality",
             "Section 8",
             "72.9% first-pass clean rate (95% CI: 58.4%–84.3%)"],
            ["RQ3 — Failure modes",
             "Section 8.3",
             "Forbidden language (38.9%), hallucination (16.7%), logic errors (11.1%)"],
            ["RQ4 — Comparison with literature",
             "Section 10",
             "Competitive with pure LLM; lower cost; structural guarantees via scaffolding"],
        ],
        col_widths=[4, 3, 6]
    )

    doc.add_page_break()

    # ===== 3. PROJECT SCALE =====
    doc.add_heading("3. Project Scale — Azure DevOps Data", level=1)

    doc.add_paragraph(
        "Data was retrieved from the Azure DevOps REST API v7.1 using WIQL queries against "
        "the Env\\ENV Kanda area path. The board contains stories across three columns: "
        "Most Wanted (backlog), Development (in progress), and Quality Assurance (testing)."
    )

    doc.add_heading("3.1 Board Overview", level=2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Total user stories on board", "55"],
            ["Stories in Quality Assurance", "42"],
            ["Stories in Development", "8"],
            ["Stories in Most Wanted (backlog)", "5"],
            ["Stories with test cases (>0)", "50"],
            ["Stories with 0 test cases", "5"],
            ["Total test cases in ADO", "870"],
        ],
        col_widths=[8, 4]
    )

    doc.add_paragraph()
    doc.add_heading("3.2 Test Cases by Board Phase", level=2)
    add_styled_table(doc,
        ["Board Column", "Stories", "Test Cases", "Avg TCs/Story"],
        [
            ["Quality Assurance", "42", "686", "16.3"],
            ["Development", "8", "138", "17.3"],
            ["Most Wanted", "5", "46", "9.2"],
            ["Total", "55", "870", "15.8"],
        ],
        col_widths=[5, 3, 3, 3]
    )

    doc.add_page_break()

    # ===== 3. AI-TEST-GEN METRICS =====
    doc.add_heading("4. AI-Test-Gen Tool — Generation Metrics", level=1)

    doc.add_paragraph(
        "Detailed generation metrics are available for 12 user stories where full local output "
        "(CSV, JSON, Objectives) was retained. These represent the most recent generation runs "
        "and provide granular traceability from acceptance criteria to individual test steps."
    )

    doc.add_heading("4.1 Aggregate Metrics (12 Stories)", level=2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Stories processed", "12"],
            ["Total acceptance criteria (ACs)", "91"],
            ["Total test cases generated", "207"],
            ["Total test steps generated", "1,563"],
            ["Accessibility test cases (WCAG 2.1 AA)", "30 (14.5%)"],
            ["Avg test cases per story", "17.2"],
            ["Avg test steps per test case", "7.6"],
            ["Avg test cases per AC", "2.3"],
            ["AC-to-TC amplification ratio", "2.3x"],
        ],
        col_widths=[8, 4]
    )

    doc.add_paragraph()
    doc.add_heading("4.2 Per-Story Breakdown", level=2)
    add_styled_table(doc,
        ["#", "Story ID", "Title", "ACs", "TCs", "Steps", "In ADO"],
        [
            ["1", "269496", "Model Space and Canvas", "4", "9", "51", "7"],
            ["2", "270457", "Zoom Controls (Keyboard & Mouse)", "7", "14", "90", "14"],
            ["3", "272575", "Live Measurement While Drawing a Line", "7", "17", "148", "15"],
            ["4", "270472", "Mirror Tool", "6", "17", "155", "26"],
            ["5", "269899", "Zoom Control", "8", "18", "127", "20"],
            ["6", "270740", "Properties Panel — Layers", "12", "32", "215", "30"],
            ["7", "271309", "File — Open/Save/Save As/Close", "8", "21", "166", "18"],
            ["8", "272261", "Show/Hide Rulers, Scale, Compass", "5", "17", "131", "10"],
            ["9", "272776", "Show/Hide Property Panels", "5", "12", "89", "13"],
            ["10", "270736", "Dimensions — Set Scale", "11", "16", "133", "16"],
            ["11", "270741", "Properties Panel — History", "9", "16", "109", "14"],
            ["12", "271916", "Print — Export to PDF", "9", "18", "149", "17"],
            ["", "Totals", "", "91", "207", "1,563", "200"],
        ],
        col_widths=[1, 2, 5.5, 1.2, 1.2, 1.5, 1.5]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Note: 'In ADO' count may differ from generated TCs due to human review "
        "(tests removed, merged, or added during QA review before upload)."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===== 4. COMPLETE INVENTORY =====
    doc.add_heading("5. Complete User Story & Test Case Inventory (ADO)", level=1)

    doc.add_paragraph(
        "The following table lists all 55 user stories from the Azure DevOps board with their "
        "current board column and test case count as retrieved via the ADO API."
    )

    # QA stories
    doc.add_heading("5.1 Quality Assurance Phase (42 stories, 686 test cases)", level=2)
    qa_stories = [
        ["269496", "Model Space and Canvas", "7"],
        ["269893", "Display Rulers on Canvas", "13"],
        ["269899", "Zoom Control", "20"],
        ["269901", "Scale Control", "10"],
        ["269903", "Coordinate Display", "10"],
        ["270033", "Compass (North Indicator)", "15"],
        ["270466", "Select / Move Tool", "16"],
        ["270472", "Mirror Tool", "26"],
        ["270479", "Basic Shape Tools (Line, Rectangle, Circle)", "66"],
        ["270480", "Pen / Pencil Tool (Hand-Draw and Connected Lines)", "18"],
        ["270542", "Hand Tool", "9"],
        ["270736", "Dimensions — Set Scale (Manual via Reference Line)", "16"],
        ["270739", "Properties Panel — Design", "11"],
        ["270740", "Properties Panel — Layers", "30"],
        ["270741", "Properties Panel — History", "14"],
        ["271053", "Error Handling — Unexpected Technical Errors & Logging", "0"],
        ["271307", "File — New Document Dialog", "13"],
        ["271309", "File — Open / Save / Save As / Close Project", "18"],
        ["271914", "Edit Menu — Basic Edit Operations", "23"],
        ["271916", "Print — Export to PDF and Physical Printer", "17"],
        ["272261", "Show / Hide Rulers, Scale, Compass", "10"],
        ["272265", "Full Screen Mode (OS native)", "10"],
        ["272279", "Tools — Settings: Canvas Display", "12"],
        ["272283", "Tools — Settings: Units & Measurement Tab", "9"],
        ["272574", "Calculate Area and Total Length", "25"],
        ["272575", "Live Measurement While Drawing a Line", "15"],
        ["272776", "Show / Hide Property Panels", "13"],
        ["272777", "Draw — Measurement Line", "19"],
        ["272778", "Dimensions — Angular", "14"],
        ["272779", "Dimensions — Radius", "7"],
        ["272780", "Dimensions — Diameter", "10"],
        ["272888", "Draw Order (Object Stacking Controls)", "18"],
        ["272889", "Object Transformation Tools (Rotate & Mirror)", "13"],
        ["272971", "Help: Contact Support", "11"],
        ["272972", "Home Screen", "14"],
        ["272973", "Help: User Manual", "9"],
        ["272981", "Help: About App", "11"],
        ["273166", "Typography Controls (Text Formatting Panel)", "15"],
        ["273167", "Show/Hide Grid (View Menu)", "9"],
        ["273566", "Canvas Label Repositioning for Improved Readability", "9"],
        ["273567", "Interface — Top Menu Toolbar (Desktop)", "64"],
        ["273568", "Interface — Left-side Toolbox (Desktop)", "17"],
    ]
    add_styled_table(doc,
        ["Story ID", "Title", "# TCs"],
        qa_stories,
        col_widths=[2.2, 9, 1.8]
    )

    doc.add_paragraph()
    doc.add_heading("5.2 Development Phase (8 stories, 138 test cases)", level=2)
    dev_stories = [
        ["270171", "Microsoft Store — Initial Setup and Configuration", "0"],
        ["270172", "Apple App Store Connect Configuration", "0"],
        ["270173", "Google Play Console Configuration", "0"],
        ["270457", "Zoom Controls (Keyboard & Mouse)", "14"],
        ["270471", "Move — Rotate Tool", "41"],
        ["270737", "GPS Input and Mapping", "37"],
        ["270738", "Undo / Redo Functionality", "35"],
        ["271917", "Insert — Images, Text, Symbols & Annotations", "11"],
    ]
    add_styled_table(doc,
        ["Story ID", "Title", "# TCs"],
        dev_stories,
        col_widths=[2.2, 9, 1.8]
    )

    doc.add_paragraph()
    doc.add_heading("5.3 Most Wanted / Backlog (5 stories, 46 test cases)", level=2)
    backlog_stories = [
        ["271482", "Format — Edit Canvas", "0"],
        ["272264", "Line Style Options — Color, Thickness, Dashed", "10"],
        ["272280", "Tools — Settings: File Management Tab", "17"],
        ["272887", "Workspace Layout (Presentation Mode)", "19"],
        ["272976", "Help: Check for Updates", "0"],
    ]
    add_styled_table(doc,
        ["Story ID", "Title", "# TCs"],
        backlog_stories,
        col_widths=[2.2, 9, 1.8]
    )

    doc.add_page_break()

    # ===== 6. TIME & PRODUCTIVITY =====
    doc.add_heading("6. Time & Productivity Analysis", level=1)

    doc.add_paragraph(
        "Time estimates are based on industry benchmarks for manual test case writing. "
        "Multiple sources indicate that a senior QA engineer typically requires 15–30 minutes "
        "to write a single detailed test case with preconditions, actions, and expected results "
        "[1][2][3]. We use 20 minutes as a conservative average for a senior QA engineer."
    )
    doc.add_paragraph(
        "Human cost baseline: $40/hr. This is a conservative estimate below the U.S. national median. "
        "According to the U.S. Bureau of Labor Statistics (BLS), the median hourly wage for "
        "Software Quality Assurance Analysts and Testers (SOC 15-1253) was $49.33/hr "
        "($102,610/year) in May 2024 [11]. Glassdoor reports $49/hr and ZipRecruiter reports "
        "$45.75/hr for QA Engineers in the U.S. as of February 2026 [12][13]. "
        "We use $40/hr to provide a conservative lower-bound estimate that accounts for "
        "variation across regions, experience levels, and company sizes. All cost savings "
        "figures in this report would be higher at the BLS median rate."
    )

    doc.add_heading("6.1 Manual Test Case Writing Benchmarks", level=2)
    add_styled_table(doc,
        ["QA Experience Level", "Time per Test Case", "Source"],
        [
            ["Junior QA Engineer", "30–45 min", "Industry surveys [1][2]"],
            ["Senior QA Engineer", "15–25 min", "Industry surveys [1][2]"],
            ["Benchmark used", "20 min (conservative avg)", "—"],
        ],
        col_widths=[5, 4, 4]
    )

    doc.add_paragraph()
    doc.add_heading("6.2 AI-Test-Gen Tool Timing", level=2)
    add_styled_table(doc,
        ["Activity", "Time", "Notes"],
        [
            ["LLM generation per story", "~2–3 min", "Gemini 2.5 Flash API call"],
            ["Total generation (12 stories)", "~30 min", "Machine time (parallelizable)"],
            ["Human QA review per story", "~15–30 min", "Quality check + fixes"],
            ["Total review (12 stories)", "~4–6 hrs", "Based on 2 reviewed stories"],
        ],
        col_widths=[5, 3, 5]
    )

    doc.add_paragraph()
    doc.add_heading("6.3 Productivity Comparison", level=2)

    # 12 stories
    p = doc.add_paragraph()
    run = p.add_run("Sampled Stories (12 stories, 207 test cases):")
    run.bold = True
    add_styled_table(doc,
        ["Approach", "Effort (hours)", "Cost (@$40/hr*)", "Savings"],
        [
            ["Fully Manual", "69.0 hrs", "$2,760", "—"],
            ["AI-Assisted (generation + review)", "~5.5 hrs", "$220 + $0.50 LLM", "—"],
            ["Time Saved", "63.5 hrs", "$2,539", "92%"],
        ],
        col_widths=[5, 3, 3.5, 2]
    )

    doc.add_paragraph()
    # Full project
    p = doc.add_paragraph()
    run = p.add_run("Full Project Projection (55 stories, 870 test cases):")
    run.bold = True
    add_styled_table(doc,
        ["Approach", "Effort (hours)", "Cost (@$40/hr*)", "Savings"],
        [
            ["Fully Manual", "290 hrs", "$11,600", "—"],
            ["AI-Assisted (generation + review)", "~23 hrs", "$920 + $2 LLM", "—"],
            ["Time Saved", "267 hrs", "$10,678", "92%"],
        ],
        col_widths=[5, 3, 3.5, 2]
    )

    doc.add_page_break()

    # ===== 7. LLM COST ANALYSIS =====
    doc.add_heading("7. LLM Cost Analysis", level=1)

    doc.add_heading("7.1 Gemini 2.5 Flash Pricing (February 2026)", level=2)
    doc.add_paragraph(
        "Pricing data retrieved from the official Google AI for Developers pricing page. "
        "The project uses Gemini 2.5 Flash as the primary LLM provider, configured via "
        "the project YAML and invoked through the google-genai SDK."
    )
    add_styled_table(doc,
        ["Tier", "Input (per 1M tokens)", "Output (per 1M tokens)"],
        [
            ["Standard", "$0.30", "$2.50"],
            ["Batch (50% discount)", "$0.15", "$1.25"],
        ],
        col_widths=[4, 4, 4]
    )

    doc.add_paragraph()
    doc.add_heading("7.2 Estimated Token Usage Per Story", level=2)
    add_styled_table(doc,
        ["Component", "Input Tokens", "Output Tokens"],
        [
            ["System prompt (rules, constraints)", "~2,500", "—"],
            ["Story context (ACs, description)", "~1,000", "—"],
            ["Reference steps (ChromaDB retrieval)", "~800", "—"],
            ["LLM correction response", "—", "~4,000"],
            ["Edge case generation", "~1,500", "~2,000"],
            ["Accessibility generation", "~1,200", "~1,500"],
            ["Total per story (est.)", "~7,000", "~7,500"],
        ],
        col_widths=[6, 3, 3]
    )

    doc.add_paragraph()
    doc.add_heading("7.3 Cost Summary", level=2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Cost per story (standard tier)", "~$0.04"],
            ["Total LLM cost (12 stories)", "~$0.48"],
            ["Projected LLM cost (55 stories)", "~$2.20"],
            ["Cost per test case", "$0.002"],
            ["Cost per test step", "$0.0003"],
        ],
        col_widths=[8, 4]
    )

    doc.add_paragraph()
    doc.add_heading("7.4 LLM vs Human Cost Comparison", level=2)
    add_styled_table(doc,
        ["Metric", "Human (Senior QA)", "AI-Test-Gen", "Cost Ratio"],
        [
            ["Cost per test case", "$13.33", "$0.002", "6,665x cheaper"],
            ["Cost for 207 TCs (12 stories)", "$2,760", "$0.48 (LLM only)", "5,750x cheaper"],
            ["Cost for 207 TCs (incl. review)", "$2,760", "$220.48", "12.5x cheaper"],
            ["Cost for 870 TCs (full project)", "$11,600", "$922 (incl. review)", "12.6x cheaper"],
        ],
        col_widths=[4.5, 3.5, 3.5, 3]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Note: 'Incl. review' accounts for human QA review time at $40/hr. "
        "Pure LLM cost is negligible ($0.002/TC); the dominant cost in AI-assisted "
        "workflow is human review, not computation.\n\n"
        "*$40/hr is a conservative estimate. The U.S. BLS median for QA Analysts "
        "(SOC 15-1253) is $49.33/hr as of May 2024 [11]. At the median rate, "
        "cost savings would be ~23% higher than reported here."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===== 8. QUALITY METRICS =====
    doc.add_heading("8. Quality Metrics", level=1)

    doc.add_paragraph(
        "Quality was assessed through manual expert review of 2 user stories (32 test cases "
        "across 3 review passes). The review applied 6 quality dimensions derived from "
        "test case design best practices and acceptance testing standards."
    )

    doc.add_heading("8.1 Sample Selection Justification", level=2)
    doc.add_paragraph(
        "Two stories were selected for detailed quality review using stratified purposive sampling "
        "to represent different levels of story complexity and test generation difficulty:"
    )
    sample_items = [
        ("Story 270741 (Properties Panel — History): ",
         "9 ACs, medium complexity, UI interaction-heavy. Represents the typical story "
         "with moderate acceptance criteria count and standard functional test patterns."),
        ("Story 271916 (Print — Export to PDF): ",
         "9 ACs, high complexity, involves external hardware interaction (printers), "
         "file system operations, and error handling. Represents the most challenging "
         "category where LLM hallucination and non-deterministic language are most likely."),
    ]
    for prefix, desc in sample_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    doc.add_paragraph(
        "Together, these stories cover 32 test cases (15.5% of the 207 generated), "
        "spanning functional, edge case, error handling, and accessibility test types. "
        "Story 271916 was reviewed twice (before and after code-level fixes) to assess "
        "the impact of prompt engineering improvements, yielding 3 review passes total."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Limitation: A sample of 2 stories (n=32 test cases) is small. While the stories were "
        "purposively selected to span difficulty levels, the findings may not generalise to all "
        "55 stories. Phase 2 will expand quality review to at least 5 stories to improve "
        "statistical power. See Section 12 (Threats to Validity) for further discussion."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_heading("8.2 First-Pass Quality Rate", level=2)
    add_styled_table(doc,
        ["Story", "TCs Reviewed", "Issues Found", "Tests Affected", "Clean Rate"],
        [
            ["270741 (Properties Panel — History)", "16", "7", "5", "68.8%"],
            ["271916 (Print/Export) — Generation 1", "16", "6", "4", "75.0%"],
            ["271916 (Print/Export) — Generation 2*", "16", "5", "4", "75.0%"],
            ["Average", "—", "—", "—", "72.9%"],
        ],
        col_widths=[5, 2.2, 2.2, 2.2, 2.2]
    )
    p = doc.add_paragraph()
    run = p.add_run("* Generation 2 ran after code-level fixes to prompt engineering and platform filtering.")
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading("8.2.1 Statistical Confidence", level=3)
    doc.add_paragraph(
        "To assess the reliability of the observed 72.9% first-pass quality rate, "
        "we compute a 95% confidence interval using the Wilson score method, which is "
        "recommended for small samples and proportions near boundary values."
    )

    # Wilson score CI calculation explanation
    # n=48 (total reviews: 16+16+16), defective=13 (5+4+4), clean=35
    # p_hat = 35/48 = 0.729
    # Wilson score 95% CI: (0.584, 0.843)
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Total test case reviews (n)", "48 (16 + 16 + 16 across 3 passes)"],
            ["Test cases with zero issues", "35"],
            ["Observed proportion (p-hat)", "0.729 (72.9%)"],
            ["95% CI (Wilson score)", "(0.584, 0.843)"],
            ["Interpretation", "We are 95% confident the true first-pass quality rate "
             "lies between 58.4% and 84.3%"],
        ],
        col_widths=[5, 8]
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "Note: The wide interval (26 percentage points) reflects the small sample size (n=48). "
        "Expanding quality review to 5+ stories in Phase 2 would narrow this interval "
        "substantially, providing stronger evidence for the observed quality rate."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading("8.3 Issue Classification", level=2)
    add_styled_table(doc,
        ["Issue Type", "Count", "% of Total", "Severity"],
        [
            ["Forbidden language ('or'/'if'/'either' in steps)", "7", "38.9%", "Critical"],
            ["Hallucinated content (invented names/data)", "3", "16.7%", "Critical"],
            ["Logical contradiction / incorrect logic", "2", "11.1%", "Major"],
            ["Missing AC coverage", "2", "11.1%", "Major"],
            ["Duplicate / overlapping tests", "2", "11.1%", "Minor"],
            ["Missing setup / precondition", "1", "5.6%", "Minor"],
            ["Inconsistent labels / terminology", "1", "5.6%", "Minor"],
            ["Total issues (across 3 reviews)", "18", "100%", "—"],
        ],
        col_widths=[6, 1.8, 2, 2.5]
    )

    doc.add_paragraph()
    doc.add_heading("8.4 Quality Dimensions", level=2)
    add_styled_table(doc,
        ["Dimension", "Description", "Pass Rate"],
        [
            ["AC Coverage", "Every AC has at least one test", "94.4%"],
            ["Deterministic Steps", "No 'or'/'if' in actions/expected results", "77.8%"],
            ["Grounded Content", "No hallucinated data not in ACs", "91.7%"],
            ["Logical Correctness", "Steps make logical sense end-to-end", "94.4%"],
            ["Consistent Setup", "Same precondition pattern across story", "66.7% → 100%*"],
            ["No Duplicates", "Each test is unique in purpose", "88.9% → 100%*"],
        ],
        col_widths=[3.5, 5.5, 3]
    )
    p = doc.add_paragraph()
    run = p.add_run("* After code-level fixes to prompt_builder.py and corrector.py.")
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading("8.5 Quality Improvement After Code Fixes", level=2)
    add_styled_table(doc,
        ["Metric", "Before Fixes", "After Fixes", "Improvement"],
        [
            ["Forbidden language violations", "3 critical + 3 medium", "1 critical + 2 medium", "50% reduction"],
            ["Setup step consistency", "Inconsistent", "Fully consistent", "Resolved"],
            ["Duplicate tests", "Present", "Eliminated", "Resolved"],
            ["Platform filtering (tablet scope)", "Incorrect", "Correct", "Resolved"],
        ],
        col_widths=[4, 3.5, 3.5, 2.5]
    )

    doc.add_page_break()

    # ===== 9. COVERAGE =====
    doc.add_heading("9. Coverage Analysis", level=1)

    doc.add_heading("9.1 Test Type Distribution (207 test cases)", level=2)
    add_styled_table(doc,
        ["Test Type", "Count", "Percentage"],
        [
            ["Functional (AC-based)", "91", "44.0%"],
            ["Edge Case / Error Handling", "56", "27.0%"],
            ["Accessibility (WCAG 2.1 AA)", "30", "14.5%"],
            ["UI Interaction / Usability", "30", "14.5%"],
            ["Total", "207", "100%"],
        ],
        col_widths=[5, 3, 3]
    )

    doc.add_paragraph()
    doc.add_heading("9.2 Platform Coverage", level=2)
    add_styled_table(doc,
        ["Platform", "Accessibility Tests", "Notes"],
        [
            ["Windows 11", "12", "All 12 stories"],
            ["iPad", "9", "9 stories (3 filtered by platform scope)"],
            ["Android Tablet", "9", "9 stories (3 filtered by platform scope)"],
        ],
        col_widths=[4, 3, 6]
    )

    doc.add_paragraph()
    doc.add_heading("9.3 AC Amplification by Story Complexity", level=2)
    add_styled_table(doc,
        ["AC Count Range", "Stories", "Avg TCs Generated", "TC:AC Ratio"],
        [
            ["4–5 ACs (simple)", "3", "12.7", "2.8x"],
            ["6–8 ACs (medium)", "6", "17.2", "2.4x"],
            ["9–12 ACs (complex)", "3", "22.0", "2.1x"],
        ],
        col_widths=[4, 2.5, 3.5, 3]
    )
    doc.add_paragraph(
        "Observation: Higher AC counts lead to slightly lower amplification ratios, "
        "suggesting the tool avoids redundancy as story complexity grows."
    )

    doc.add_page_break()

    # ===== 10. COMPARISON WITH RELATED WORK =====
    doc.add_heading("10. Comparison with Related Work", level=1)

    doc.add_paragraph(
        "The following table positions our Phase 1 results against findings from "
        "recent peer-reviewed studies on LLM-based test case generation."
    )

    add_styled_table(doc,
        ["Study", "Approach", "Key Metric", "Our Result"],
        [
            [
                "Li & Yuan (2024) [4]\nBeihang University",
                "Pure LLM\n(GPT-4, 0-shot)",
                "Accuracy: 84.6% (easy)\n57.9% (hard)",
                "First-pass clean: 72.9%\n(functional tests)"
            ],
            [
                "Masuda et al. (2025) [5]\nIndustry Study",
                "LLM prompt-based\n(no RAG)",
                "Macro-recall: 0.81\n(Bluetooth spec)",
                "AC coverage: 94.4%"
            ],
            [
                "Xue et al. (2024) [6]\nLLM4Fin (ISSTA)",
                "Hybrid: Fine-tuned\nLLM + Algorithm",
                "Coverage: 98.18%\nTime: 20 min → 7 sec",
                "Coverage: 94.4%\nTime: 69 hrs → 5.5 hrs"
            ],
            [
                "Our approach",
                "Hybrid: Rule-based\n+ LLM correction",
                "Cost: $0.002/TC\nTime: 92% reduction",
                "870 TCs across\n55 stories"
            ],
        ],
        col_widths=[3.5, 3, 3.5, 3.5]
    )

    doc.add_paragraph()
    doc.add_heading("10.1 Positioning", level=2)
    doc.add_paragraph(
        "Our hybrid approach differs from pure LLM studies [4][5] by using deterministic "
        "rule-based scaffolding before LLM correction, which provides structural consistency "
        "and reduces hallucination. Compared to LLM4Fin [6], which fine-tunes models for "
        "domain-specific extraction, our approach uses prompt engineering with retrieval-augmented "
        "reference steps (ChromaDB), making it more generalizable across projects without "
        "model retraining."
    )
    doc.add_paragraph(
        "The systematic review by Camara & Fonseca (2025) [7] categorized 84 studies into "
        "four approaches: prompt engineering, feedback-driven, fine-tuning, and hybrid. "
        "Our work falls into the hybrid category, combining prompt engineering with "
        "rule-based preprocessing — a combination with limited prior research."
    )

    doc.add_page_break()

    # ===== 11. KPIs =====
    doc.add_heading("11. Key Performance Indicators (KPIs)", level=1)

    add_styled_table(doc,
        ["KPI", "Value", "Thesis Target (O5)", "Status"],
        [
            ["Time reduction vs manual", "92%", "≥ 60%", "EXCEEDED"],
            ["LLM cost per test case", "$0.002", "< $1.00", "EXCEEDED"],
            ["First-pass quality rate", "72.9%", "≥ 80%", "BELOW TARGET"],
            ["AC coverage rate", "94.4%", "100%", "CLOSE"],
            ["TC-to-AC amplification", "2.3x", "≥ 2.0x", "MET"],
            ["Accessibility test inclusion", "14.5%", "> 0%", "MET"],
            ["Platform filtering accuracy", "100% (after fix)", "100%", "MET"],
        ],
        col_widths=[4, 2.5, 3, 3]
    )

    doc.add_paragraph()
    doc.add_heading("11.1 Strengths", level=2)
    strengths = [
        "Massive time and cost savings (92% reduction, 5,750x cheaper for LLM computation alone)",
        "Strong AC coverage (94.4%) with 2.3x amplification from ACs to test cases",
        "Consistent structure and setup patterns across generated tests (after code fixes)",
        "Automatic accessibility test inclusion for all supported platforms",
        "Project-agnostic YAML configuration — adaptable to different applications",
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("11.2 Areas for Improvement", level=2)
    improvements = [
        "First-pass quality (72.9%) is below the 80% target — error/edge case tests are the weakest category",
        "Forbidden language ('or'/'if') persists in ~23% of error-handling tests despite prompt rules",
        "LLM occasionally hallucinates specific names (e.g., printer models) not grounded in ACs",
        "Human review still required (~25 min/story) — not yet fully autonomous",
    ]
    for i in improvements:
        doc.add_paragraph(i, style='List Bullet')

    doc.add_page_break()

    # ===== 12. THREATS TO VALIDITY =====
    doc.add_heading("12. Threats to Validity", level=1)

    doc.add_paragraph(
        "This section discusses potential threats to the validity of the empirical findings "
        "presented in this report, following the classification framework of Wohlin et al. (2012)."
    )

    doc.add_heading("12.1 Internal Validity", level=2)
    internal_threats = [
        ("Single-reviewer bias: ",
         "All quality reviews were conducted by the same researcher who also developed the tool. "
         "This introduces potential confirmation bias — the reviewer may unconsciously apply "
         "lenient judgment to output from their own system. Mitigation: A structured 6-dimension "
         "rubric with explicit pass/fail criteria was applied consistently across all reviews. "
         "Phase 2 should include an independent second reviewer for inter-rater reliability."),
        ("Learning effect in iterative review: ",
         "Story 271916 was reviewed twice (before and after code fixes), meaning the reviewer "
         "had prior knowledge of expected issues in the second pass. This may inflate the "
         "perceived improvement. Mitigation: Both passes used the same rubric, and the "
         "same issues were tracked by type and location for objective comparison."),
        ("Time estimation baseline: ",
         "The 20-minute-per-test-case manual baseline is derived from industry surveys [1][2][3] "
         "rather than from direct measurement of the team's actual writing speed. Actual times "
         "may vary by team experience and story complexity. The $40/hr rate is conservative "
         "relative to BLS median ($49.33/hr) [11]."),
    ]
    for prefix, desc in internal_threats:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)

    doc.add_heading("12.2 External Validity", level=2)
    external_threats = [
        ("Single-project domain: ",
         "All evaluation was conducted on one application (ENV QuickDraw — a CAD-style drawing tool). "
         "The tool's effectiveness may differ for other domains (e.g., e-commerce, fintech, mobile apps) "
         "where acceptance criteria follow different patterns. The YAML-based project configuration "
         "is designed for portability, but cross-domain validation is needed."),
        ("Single LLM provider: ",
         "All generation used Gemini 2.5 Flash. Performance, cost, and quality metrics may differ "
         "with GPT-4o, Claude Sonnet, or open-source models. Phase 2 will include a formal "
         "multi-provider comparison (RQ4)."),
        ("English-only requirements: ",
         "All user stories and acceptance criteria were written in English. The tool has not been "
         "evaluated on requirements in other languages."),
    ]
    for prefix, desc in external_threats:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)

    doc.add_heading("12.3 Construct Validity", level=2)
    construct_threats = [
        ("Quality dimension selection: ",
         "The 6 quality dimensions (AC coverage, determinism, grounding, logic, consistency, "
         "uniqueness) were derived from test case design best practices and IEEE 829 standards, "
         "but they may not capture all aspects of test case quality relevant to practitioners. "
         "For instance, test executability and maintainability are not directly measured."),
        ("First-pass quality as proxy: ",
         "The 'first-pass quality rate' measures structural correctness, not functional "
         "effectiveness (i.e., whether the test cases actually find bugs in the application). "
         "A test case can be structurally clean but fail to detect real defects. "
         "Functional effectiveness evaluation is deferred to Phase 2."),
    ]
    for prefix, desc in construct_threats:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)

    doc.add_heading("12.4 Conclusion Validity", level=2)
    conclusion_threats = [
        ("Small quality review sample: ",
         "Quality metrics are based on 32 test cases across 2 stories (3 review passes, n=48). "
         "The resulting 95% confidence interval for the first-pass quality rate is wide "
         "(58.4%–84.3%), limiting the precision of the estimate. Expanding to 5+ stories "
         "in Phase 2 will improve statistical power."),
        ("No formal hypothesis testing: ",
         "This report presents descriptive statistics and comparative benchmarks rather than "
         "formal inferential tests (e.g., paired t-test for manual vs. AI-assisted time). "
         "Formal statistical testing requires larger sample sizes and controlled conditions, "
         "planned for Phase 2."),
    ]
    for prefix, desc in conclusion_threats:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ===== 13. REFERENCES =====
    doc.add_heading("13. References", level=1)

    references = [
        "[1] Apriorit, \"Techniques for Time Estimation in Software Testing,\" 2024. "
        "Available: https://www.apriorit.com/qa-blog/197-testing-time-estimation",

        "[2] SoftwareTestingHelp, \"Software Test Estimation Techniques,\" 2024. "
        "Available: https://www.softwaretestinghelp.com/software-test-estimation-how-to-estimate-testing-time-accurately/",

        "[3] Idea Link, \"Software Testing Costs: 20–40% of Development Budget,\" 2024. "
        "Available: https://idealink.tech/blog/understanding-software-testing-costs-development-breakdown",

        "[4] K. Li and Y. Yuan, \"Large Language Models as Test Case Generators: "
        "Performance Evaluation and Enhancement,\" Beihang University, 2024. "
        "arXiv:2404.13340",

        "[5] S. Masuda, S. Kouzawa, K. Sezai, H. Suhara, Y. Hiruta, and K. Kudou, "
        "\"Generating High-Level Test Cases from Requirements using LLM: An Industry Study,\" "
        "2025. arXiv:2510.03641",

        "[6] Z. Xue, L. Li, S. Tian, X. Chen, P. Li, L. Chen, T. Jiang, and M. Zhang, "
        "\"LLM4Fin: Fully Automating LLM-Powered Test Case Generation for FinTech Software "
        "Acceptance Testing,\" in Proc. ISSTA 2024, ACM, 2024.",

        "[7] R. Camara and C. Fonseca, \"A Review of Large Language Models for Automated "
        "Test Case Generation,\" Machine Learning and Knowledge Extraction, vol. 7, no. 3, "
        "p. 97, MDPI, 2025.",

        "[8] ACM ISEC 2025, \"Test Case Generation for Requirements in Natural Language — "
        "An LLM Comparison Study,\" in Proc. 18th Innovations in Software Engineering "
        "Conference, ACM, 2025. doi:10.1145/3717383.3717389",

        "[9] Google, \"Gemini Developer API Pricing,\" February 2026. "
        "Available: https://ai.google.dev/gemini-api/docs/pricing",

        "[10] TestRail, \"Guide to the Top 20 QA Metrics That Matter,\" 2024. "
        "Available: https://www.testrail.com/blog/qa-metrics-matter/",

        "[11] U.S. Bureau of Labor Statistics, \"Occupational Employment and Wages, "
        "May 2024: 15-1253 Software Quality Assurance Analysts and Testers.\" "
        "Median hourly wage: $49.33; Median annual wage: $102,610. "
        "Available: https://www.bls.gov/oes/current/oes151253.htm",

        "[12] Glassdoor, \"QA Engineer Salaries in the United States,\" February 2026. "
        "Average: $101,379/year (~$49/hr). "
        "Available: https://www.glassdoor.com/Salaries/qa-engineer-salary-SRCH_KO0,11.htm",

        "[13] ZipRecruiter, \"QA Engineer Salary,\" February 2026. "
        "Average: $95,168/year (~$45.75/hr). "
        "Available: https://www.ziprecruiter.com/Salaries/Qa-Engineer-Salary",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(9)

    # ===== SAVE =====
    output_dir = Path("/Users/gulzhasmailybayeva/Desktop/test_gen/output")
    output_path = output_dir / "PHASE1_EMPIRICAL_METRICS_REPORT.docx"
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
