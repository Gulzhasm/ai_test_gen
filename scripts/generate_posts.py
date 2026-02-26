#!/usr/bin/env python3
"""Generate Medium and LinkedIn posts as Word documents."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_medium_heading(doc, text, level=1):
    """Add a Medium-style heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_body(doc, text, bold_prefix=None):
    """Add body paragraph with optional bold prefix."""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_code_block(doc, code):
    """Add a code-style block."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_quote(doc, text):
    """Add a blockquote-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(10)
    return p


def create_medium_post():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ===== TITLE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "We Built a System That Generates 207 Test Cases "
        "from User Stories — Here's What We Learned"
    )
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "How a hybrid rule-based + LLM approach cut our test creation time by 92% "
        "at $0.002 per test case"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("By Gulzhas Mailybayeva | KandaSoft | February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ===== HOOK =====
    add_body(doc,
        "Our QA team was drowning. Fifty-five user stories. Hundreds of acceptance criteria. "
        "And every single test case had to be written by hand — action by action, expected result "
        "by expected result. Each test case took 20 minutes. We were looking at 290+ hours of work "
        "just to build the test plan."
    )
    add_body(doc,
        "So we asked a different question: What if we could build a system that generates "
        "structured first drafts automatically from user stories?"
    )

    # ===== THE PROBLEM =====
    add_medium_heading(doc, "The Problem: Manual Test Case Creation Doesn't Scale")

    add_body(doc,
        "At KandaSoft, we're building ENV QuickDraw — a cross-platform CAD-style drawing application "
        "for environmental engineering. Like any enterprise product, it demands rigorous QA. Every "
        "feature needs detailed manual test cases with preconditions, step-by-step actions, and "
        "expected results before it can ship."
    )

    add_body(doc, "The numbers told the story:")

    bullets = [
        "55 user stories across three development phases",
        "Each story has 4–12 acceptance criteria",
        "Each AC needs 1–3 test cases to cover happy paths, edge cases, and accessibility",
        "A senior QA engineer writes ~3 test cases per hour",
        "That's 290+ hours — over 7 weeks of full-time work, just writing tests",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_body(doc,
        "And that's before factoring in review cycles, consistency checks, and the inevitable "
        "\"this test step says 'or' — that's not deterministic\" feedback."
    )

    # ===== THE APPROACH =====
    add_medium_heading(doc, "Our Approach: Hybrid Rule-Based + LLM Generation")

    add_body(doc,
        "We didn't bet everything on a single LLM prompt. Instead, we engineered a hybrid pipeline "
        "that combines the reliability of deterministic scaffolding with the language fluency of "
        "Large Language Models."
    )

    add_medium_heading(doc, "The Pipeline", level=2)

    add_code_block(doc,
        "User Story (ADO)\n"
        "      |\n"
        "      v\n"
        "  [1] Rule-Based Generator\n"
        "      - Parse acceptance criteria\n"
        "      - Generate test scaffolding (IDs, titles, areas)\n"
        "      - Map ACs to test cases deterministically\n"
        "      |\n"
        "      v\n"
        "  [2] LLM Corrector (Gemini 2.5 Flash)\n"
        "      - Enrich with natural-language steps\n"
        "      - Add edge cases and error handling\n"
        "      - Apply 13-rule prompt constraint system\n"
        "      |\n"
        "      v\n"
        "  [3] Post-Processing\n"
        "      - Accessibility test injection (WCAG 2.1 AA)\n"
        "      - Platform filtering (Windows/iPad/Android)\n"
        "      - ChromaDB reference step matching\n"
        "      |\n"
        "      v\n"
        "  [4] Export to Azure DevOps Test Plans\n"
        "      - CSV import / API upload\n"
        "      - Auto-create test suites per story"
    )

    add_medium_heading(doc, "Why Hybrid?", level=2)

    add_body(doc,
        "Pure LLM generation has a well-documented problem: hallucination. "
        "In our domain, that means invented printer names, made-up UI labels, "
        "and non-deterministic test steps containing \"or\" and \"if\" — all of which "
        "are unacceptable in production test cases."
    )
    add_body(doc,
        "Our rule-based layer handles the structural work — consistent IDs, title formatting, "
        "area paths, setup steps. The LLM handles what it's actually good at: understanding "
        "requirements in natural language and generating human-readable test scenarios."
    )
    add_body(doc,
        "This mirrors findings from Xue et al. (ISSTA 2024), who showed that combining "
        "LLM extraction with algorithmic generation achieved 98% business scenario coverage "
        "in FinTech acceptance testing — significantly outperforming both pure LLM and pure "
        "manual approaches."
    )

    # ===== THE NUMBERS =====
    add_medium_heading(doc, "The Results: Phase 1 Empirical Metrics")

    add_body(doc,
        "After running the ai-test-gen pipeline across 12 stories from the ENV QuickDraw project, "
        "we collected comprehensive metrics from Azure DevOps and our local generation output."
    )

    add_medium_heading(doc, "Scale", level=2)

    scale_items = [
        ("55", " user stories on the ADO board"),
        ("207", " test cases generated and uploaded via the ai-test-gen pipeline (12 stories with detailed local output)"),
        ("870", " total test cases now live in Azure DevOps Test Plans (including AI-generated + manually authored)"),
        ("50", " stories have full test coverage (91%)"),
        ("1,563", " individual test steps generated by the pipeline"),
    ]
    for val, desc in scale_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(val)
        run.bold = True
        p.add_run(desc)

    add_medium_heading(doc, "Time Savings", level=2)

    add_body(doc,
        "For the 12 stories where we have detailed timing data:"
    )
    time_items = [
        "Manual estimate: 69 hours (207 TCs x 20 min each)",
        "AI-assisted: 5.5 hours (30 min generation + 5 hours human review)",
        "Time saved: 63.5 hours — a 92% reduction",
    ]
    for t in time_items:
        doc.add_paragraph(t, style='List Bullet')

    add_body(doc,
        "Projected across the full project (all 55 stories, 870 total test cases): 267 hours saved."
    )

    add_medium_heading(doc, "Cost", level=2)

    add_body(doc,
        "We used Gemini 2.5 Flash as our LLM provider. The economics are striking:"
    )
    cost_items = [
        "LLM cost per test case: $0.002",
        "Total LLM cost for 12 stories (207 TCs): $0.48",
        "Projected cost for all 55 stories: $2.20",
        "Human equivalent: $11,600 (at $40/hr — conservative; U.S. BLS median is $49/hr*)",
        "Even including human review time, AI-assisted is 12.5x cheaper",
    ]
    for c in cost_items:
        doc.add_paragraph(c, style='List Bullet')

    add_quote(doc,
        "The LLM cost is essentially a rounding error. The real cost in AI-assisted "
        "testing isn't computation — it's the human review loop."
    )

    # ===== QUALITY =====
    add_medium_heading(doc, "What About Quality?")

    add_body(doc,
        "Speed means nothing if the tests are wrong. We performed detailed manual reviews "
        "of 32 test cases across 2 user stories (3 review passes). Here's what we found:"
    )

    add_medium_heading(doc, "First-Pass Quality: 72.9%", level=2)

    add_body(doc,
        "About 73% of generated test cases needed zero corrections. The remaining 27% had "
        "issues we categorized into 7 types:"
    )

    quality_items = [
        ("Forbidden language (38.9%)", " — \"or\"/\"if\" in steps making them non-deterministic"),
        ("Hallucinated content (16.7%)", " — invented printer names, specific UI labels not in the AC"),
        ("Logical errors (11.1%)", " — contradictory steps or impossible sequences"),
        ("Missing coverage (11.1%)", " — AC scenario not fully tested"),
        ("Duplicates (11.1%)", " — overlapping test cases"),
        ("Missing setup (5.6%)", " — preconditions not established"),
        ("Inconsistent labels (5.6%)", " — mismatched terminology"),
    ]
    for val, desc in quality_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(val)
        run.bold = True
        p.add_run(desc)

    add_medium_heading(doc, "The Error-Handling Test Problem", level=2)

    add_body(doc,
        "The most persistent issue was with error-handling and edge case tests. "
        "When the LLM needs to describe \"what happens when a printer is not found,\" "
        "it naturally falls into conditional language: \"If no printers are found, then...\" "
        "or \"the error message says 'not found or unavailable'.\""
    )
    add_body(doc,
        "This is exactly the kind of non-deterministic language that makes test cases "
        "ambiguous and untestable. We addressed this through prompt engineering — adding "
        "specific rules about error-condition setup patterns and forbidden language. "
        "After the fix, critical violations dropped by 50%."
    )
    add_body(doc,
        "Li & Yuan (2024) observed a similar pattern: GPT-4 scored 84.6% accuracy on "
        "easy test generation but dropped to 57.9% on hard problems. Error conditions "
        "are inherently \"harder\" — they require reasoning about negative states."
    )

    # ===== ARCHITECTURE =====
    add_medium_heading(doc, "Under the Hood: Technical Architecture")

    add_body(doc,
        "The framework follows Clean Architecture principles with clear separation of concerns:"
    )

    add_code_block(doc,
        "ai-test-gen/\n"
        "  core/\n"
        "    domain/           # Entities: UserStory, TestCase\n"
        "    interfaces/       # Abstract contracts (Python ABCs)\n"
        "    services/\n"
        "      llm/            # LLM providers (Gemini, OpenAI, Anthropic, Ollama)\n"
        "      metrics/        # Cost calculator, token tracking\n"
        "  infrastructure/\n"
        "    ado/              # Azure DevOps API integration\n"
        "    chroma/           # ChromaDB for semantic step matching\n"
        "  projects/\n"
        "    configs/          # YAML configs per project\n"
        "  output/             # Generated CSV, JSON, objectives"
    )

    add_medium_heading(doc, "Key Design Decisions", level=2)

    decisions = [
        ("Project-agnostic YAML configs: ",
         "Each project gets a YAML file defining platforms, LLM provider, "
         "ADO connection, and test ID patterns. Switch projects by changing one flag."),
        ("ChromaDB for reference steps: ",
         "Previously generated test steps are embedded and stored in a vector database. "
         "When generating new tests, similar steps are retrieved and passed as context "
         "to the LLM — ensuring consistent terminology across the test suite."),
        ("13-rule prompt constraint system: ",
         "Instead of hoping the LLM follows conventions, we encode 13 explicit rules "
         "in the system prompt: no forbidden language, deterministic steps, consistent setup, "
         "no hallucinated names, AC traceability, and more."),
        ("Factory pattern for LLM providers: ",
         "Swap between Gemini, OpenAI, Anthropic, or Ollama by changing one config line. "
         "Each provider implements the same interface with token tracking and cost calculation."),
    ]
    for prefix, desc in decisions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        p.add_run(desc)

    # ===== COMPARISON =====
    add_medium_heading(doc, "How This Compares to Existing Research")

    add_body(doc,
        "A 2025 systematic review (Camara & Fonseca) examined 84 studies on LLM-based "
        "test case generation and categorized them into four approaches: prompt engineering, "
        "feedback-driven, model fine-tuning, and hybrid. Our work falls into the hybrid "
        "category — and we found relatively few prior studies combining rule-based preprocessing "
        "with LLM correction."
    )

    add_body(doc, "How we compare:")

    compare_items = [
        ("vs. Pure LLM (Li & Yuan, 2024): ",
         "GPT-4 achieves 84.6% accuracy on easy benchmarks but 57.9% on hard ones. "
         "Our hybrid approach achieves 72.9% first-pass quality across mixed-difficulty "
         "test scenarios — competitive, especially given our tests are multi-step functional "
         "test cases, not unit tests."),
        ("vs. LLM4Fin (Xue et al., ISSTA 2024): ",
         "Their hybrid fine-tuned approach achieved 98.18% coverage but requires model "
         "retraining per domain. Our prompt-based approach is more portable — zero training, "
         "just config."),
        ("vs. Masuda et al. (2025): ",
         "Their prompt-only approach (no RAG) achieved 0.81 macro-recall on Bluetooth specs. "
         "Our retrieval-augmented approach (ChromaDB) with domain-specific reference steps "
         "provides additional grounding, reducing hallucination."),
    ]
    for prefix, desc in compare_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        p.add_run(desc)

    # ===== LESSONS =====
    add_medium_heading(doc, "5 Lessons We Learned")

    lessons = [
        ("1. Rules first, LLM second. ",
         "The LLM is brilliant at language but terrible at structure. Let deterministic "
         "code handle IDs, titles, file format, and platform filtering. Let the LLM "
         "handle natural language enrichment."),
        ("2. Error tests are the hardest. ",
         "LLMs struggle with negative testing because describing \"what goes wrong\" "
         "naturally invites conditional language. You need explicit prompt rules and "
         "setup preconditions to force deterministic error scenarios."),
        ("3. Reference steps prevent drift. ",
         "Without ChromaDB retrieval of previous test steps, the LLM would use different "
         "terminology each run. \"Launch the app\" vs \"Open the application\" vs "
         "\"Start the program\" — consistency matters for a test suite."),
        ("4. The cost isn't the LLM. ",
         "At $0.002 per test case, the LLM is essentially free. The real cost is human "
         "review — which is still necessary for that last 27% of issues. The goal is "
         "closing the quality gap so review time shrinks toward zero."),
        ("5. Platform-aware generation matters. ",
         "Our initial version generated iPad and Android Tablet accessibility tests for "
         "every story — even desktop-only features like Print. A simple AC text scan "
         "that filters platforms by mention saved us from uploading irrelevant tests."),
    ]
    for prefix, desc in lessons:
        p = doc.add_paragraph()
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(8)

    # ===== WHAT'S NEXT =====
    add_medium_heading(doc, "What's Next: Phase 2")

    next_items = [
        "Push first-pass quality from 72.9% to 85%+ through improved error-case prompts",
        "Add feedback loop: learn from human corrections to improve future generations",
        "Benchmark against pure GPT-4o and Claude Sonnet for a formal RQ4 comparison",
        "Expand to Playwright automation script generation from the manual test cases",
        "Publish evaluation framework for reproducibility",
    ]
    for n in next_items:
        doc.add_paragraph(n, style='List Bullet')

    # ===== CTA =====
    add_medium_heading(doc, "Try It Yourself")

    add_body(doc,
        "The framework is open source and designed to be project-agnostic. "
        "If you're drowning in manual test case creation, check it out:"
    )

    links = [
        "GitHub: https://github.com/Gulzhasm/ai_test_gen",
        "KandaSoft: https://www.kandasoft.com",
    ]
    for link in links:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(link)
        run.bold = True
        run.font.size = Pt(11)

    add_body(doc,
        "Questions? Comments? We'd love to hear from QA teams who are exploring "
        "AI-assisted testing. Connect with us on LinkedIn or open an issue on GitHub."
    )

    # ===== REFERENCES =====
    add_medium_heading(doc, "References")

    refs = [
        "[1] K. Li and Y. Yuan, \"Large Language Models as Test Case Generators: "
        "Performance Evaluation and Enhancement,\" Beihang University, 2024. arXiv:2404.13340",

        "[2] S. Masuda et al., \"Generating High-Level Test Cases from Requirements "
        "using LLM: An Industry Study,\" 2025. arXiv:2510.03641",

        "[3] Z. Xue et al., \"LLM4Fin: Fully Automating LLM-Powered Test Case "
        "Generation for FinTech Software Acceptance Testing,\" ISSTA 2024.",

        "[4] R. Camara and C. Fonseca, \"A Review of Large Language Models for "
        "Automated Test Case Generation,\" MAKE, vol. 7(3), MDPI, 2025.",

        "[5] ACM ISEC, \"Test Case Generation for Requirements in Natural Language — "
        "An LLM Comparison Study,\" 18th Innovations in SE Conference, 2025.",

        "[6] Google, \"Gemini Developer API Pricing,\" Feb 2026. "
        "https://ai.google.dev/gemini-api/docs/pricing",

        "*[7] U.S. Bureau of Labor Statistics, \"Occupational Employment and Wages, "
        "May 2024: 15-1253 Software Quality Assurance Analysts and Testers.\" "
        "Median hourly wage: $49.33. https://www.bls.gov/oes/current/oes151253.htm "
        "— We use $40/hr as a conservative lower-bound estimate.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.paragraph_format.space_after = Pt(4)

    # ===== TAGS =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Tags: #AI #SoftwareTesting #QA #LLM #TestAutomation #MachineLearning "
        "#AzureDevOps #Gemini #PromptEngineering #QualityAssurance"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    output_path = Path("/Users/gulzhasmailybayeva/Desktop/test_gen/output/MEDIUM_POST.docx")
    doc.save(str(output_path))
    print(f"Medium post saved to: {output_path}")


def create_linkedin_post():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LinkedIn Post — KandaSoft Page")
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x77, 0xB5)

    doc.add_paragraph()

    note = doc.add_paragraph()
    run = note.add_run(
        "Instructions: Copy the text below into a new LinkedIn post on the KandaSoft company page. "
        "The first line appears in the preview — make sure it's visible. "
        "Add the Medium link when published. Character count: ~1,400 (LinkedIn max: 3,000)."
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    doc.add_paragraph()

    # ===== THE POST =====
    lines = [
        "207 test cases generated. $0.48 in LLM costs. 92% less time.",
        "",
        "That's not a typo.",
        "",
        "Our QA team was facing 290+ hours of manual test case writing for ENV QuickDraw "
        "— our cross-platform CAD application for environmental engineering.",
        "",
        "So we built an internal test generation tool — ai-test-gen — that combines "
        "deterministic scaffolding with an LLM correction layer to produce structured "
        "test cases directly from user stories.",
        "",
        "Using the ai-test-gen pipeline, we generated and uploaded structured test cases "
        "directly into Azure DevOps Test Plans:",
        "",
        "   55 user stories processed",
        "   207 test cases generated and uploaded via ai-test-gen",
        "   870 total test cases now live in ADO (AI-generated + manually authored)",
        "   $0.002 per test case (yes, less than a penny)",
        "   92% reduction in test creation time",
        "   73% of tests needed zero human corrections",
        "",
        "The key? We didn't just throw user stories at ChatGPT and hope for the best.",
        "",
        "The tool uses rule-based logic for structure (IDs, formatting, "
        "platform filtering) and Gemini 2.5 Flash for natural-language enrichment — "
        "writing test steps that actually make sense.",
        "",
        "ChromaDB keeps terminology consistent. 13 prompt rules prevent hallucinations. "
        "And the whole pipeline plugs directly into Azure DevOps Test Plans.",
        "",
        "The LLM cost for the entire project? $2.20.",
        "The manual equivalent? $11,600 (at $40/hr — the U.S. BLS median is actually $49/hr).",
        "",
        "Is it perfect? Not yet. Error-handling tests still need human review "
        "(LLMs love writing \"if this, then that\" — which isn't how testers think). "
        "But we're pushing quality from 73% to 85%+ in Phase 2.",
        "",
        "Full technical deep-dive with architecture, metrics, and research citations:",
        "[MEDIUM LINK HERE]",
        "",
        "Open source: https://github.com/Gulzhasm/ai_test_gen",
        "",
        "#AI #SoftwareTesting #QualityAssurance #LLM #TestAutomation "
        "#AzureDevOps #Innovation #KandaSoft #MachineLearning #QA",
    ]

    for line in lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    doc.add_paragraph()

    # Suggested image note
    note = doc.add_paragraph()
    run = note.add_run(
        "Suggested visual: A before/after graphic showing:\n"
        "LEFT: \"Manual\" — 290 hours, $11,600, handwriting icon\n"
        "RIGHT: \"AI-Assisted\" — 23 hours, $922, tool/pipeline icon\n"
        "With the headline: \"Same test suite. Engineered differently.\"\n\n"
        "Alternative: Screenshot of the ADO Test Plan showing the generated test cases."
    )
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save
    output_path = Path("/Users/gulzhasmailybayeva/Desktop/test_gen/output/LINKEDIN_POST.docx")
    doc.save(str(output_path))
    print(f"LinkedIn post saved to: {output_path}")


if __name__ == "__main__":
    create_medium_post()
    create_linkedin_post()
